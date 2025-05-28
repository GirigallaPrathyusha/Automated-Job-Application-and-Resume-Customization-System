import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from supabase import create_client
import requests
import time
import re

# Configuration (same as before)
SUPABASE_URL = "https://cjftrualbdceboadyieg.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImNqZnRydWFsYmRjZWJvYWR5aWVnIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDQ5NzY0MzIsImV4cCI6MjA2MDU1MjQzMn0.c3WmjKL-y4TNAO-YtSyMUutwrzq_Un-GvrtV9UpygME"
GEMINI_API_KEY = "AIzaSyCRRo887rXeSZd5nUwrpicydXSIyQWjFw8"
MODEL_NAME = "models/gemini-1.5-pro-latest"

# Initialize clients
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)

def extract_pdf_structure(pdf_bytes):
    """Improved PDF extraction that preserves structure"""
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        blocks = []
        for page in doc:
            for block in page.get_text("blocks"):
                x0, y0, x1, y1, text, block_no, block_type = block
                if text.strip():
                    blocks.append({
                        'text': text.strip(),
                        'y_pos': y0,
                        'is_heading': block_type == 1 or (len(text) < 30 and text.isupper())
                    })
        
        # Sort by vertical position to maintain reading order
        blocks.sort(key=lambda b: b['y_pos'])
        return [b['text'] for b in blocks]

def preserve_structure_customize(original_text, job_desc):
    """Customization that strictly preserves structure"""
    prompt = f"""
    CUSTOMIZE THIS RESUME FOR THE JOB BELOW WHILE:
    1. PRESERVING THE EXACT ORIGINAL STRUCTURE
    2. KEEPING ALL SECTION HEADERS IN PLACE
    3. MAINTAINING ORIGINAL SPACING AND FORMAT
    4. ONLY MODIFYING TEXT CONTENT WITHIN EXISTING SECTIONS
    5. ADD ABOUT BASED ON JOB DESCRIPTION
    
    JOB DESCRIPTION:
    {job_desc[:5000]}
    
    RESUME (MAKE MINIMAL CHANGES TO THIS EXACT STRUCTURE):
    {original_text[:15000]}
    
    INSTRUCTIONS:
    1. Only modify text content within existing sections 
    2. For "TECH SKILLS", bold MAX 3 items matching the job
    3. For "PROFILE", rewrite to better match the job (keep same length)
    4. Never change:
       - Section headers ("GET IN TOUCH!", "EDUCATION", etc.)
       - Contact info formatting
       - Bullet points/spacing
    
    RETURN THE RESUME WITH YOUR MINIMAL CHANGES IN THE EXACT SAME STRUCTURE.
    """
    
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompt,
            generation_config={"temperature": 0.2}
        )
        return response.text
    except Exception as e:
        print(f"⚠ Customization failed: {str(e)}")
        return original_text

def rebuild_docx(structured_text):
    """Reconstruct DOCX with proper formatting"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for item in structured_text.split('\n'):
        if not item.strip():
            doc.add_paragraph()  # Preserve empty lines
            continue
            
        # Detect section headers
        if (item.isupper() and len(item) < 50) or ':' in item:
            p = doc.add_paragraph(item)
            p.style = 'Heading 2'
            p.runs[0].bold = True
        # Detect bullet points
        elif item.strip().startswith('-'):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item[1:].strip())
        else:
            doc.add_paragraph(item)
    
    return doc

def process_resumes():
    try:
        resumes = supabase.table("resumes").select("*").execute().data
        jobs = supabase.table("jobs").select("*").execute().data
        
        for resume in resumes:
            if not resume.get('file_url'):
                continue
                
            try:
                # Download resume
                response = requests.get(resume['file_url'], timeout=10)
                file_bytes = io.BytesIO(response.content)
                
                # Extract text with structure
                if resume['file_url'].lower().endswith('.docx'):
                    doc = Document(file_bytes)
                    original_text = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    original_text = '\n'.join(extract_pdf_structure(file_bytes.getvalue()))
                
                for job in jobs:
                    print(f"🔧 Enhancing resume {resume['id'][:8]} for job {job['id'][:8]}")
                    
                    # Customize while preserving structure
                    customized_text = preserve_structure_customize(original_text, job['description'])
                    
                    # Rebuild document
                    new_doc = rebuild_docx(customized_text)
                    
                    # Upload
                    output = io.BytesIO()
                    new_doc.save(output)
                    supabase.storage.from_("customizedresumes").upload(
                        f"enhanced_{resume['id']}_{job['id']}.docx",
                        output.getvalue(),
                        {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                    )
                    print(f"✅ Structured resume uploaded")
                    
                    time.sleep(1.1)
            
            except Exception as e:
                print(f"⚠ Failed to process resume {resume['id']}: {str(e)}")
                continue
                
    except Exception as e:
        print(f"💥 Fatal error: {str(e)}")

if __name__ == "__main__":
    process_resumes()